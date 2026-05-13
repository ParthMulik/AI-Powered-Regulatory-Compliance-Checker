# modifier.py
import streamlit as st
from docx import Document
from docx.shared import Pt
from docx2pdf import convert
import tempfile
import os
import pythoncom

# Assuming this helper exists in your project
from llm_helper import call_llm_with_fallback


def _rewrite_clause_with_llm(item):
    """
    Call LLM to rewrite a clause safely.
    Falls back if LLM fails.
    """
    clause = item.get("clause", "")
    prompt = (
        "Rewrite the following high-risk contract clause into a safer, "
        "more compliant version while preserving its intent:\n\n"
        f"{clause}"
    )
    return call_llm_with_fallback(prompt)


def modify_contract_docx(analysis_results):
    """
    Create a NEW Word document with:
    - Each High Risk Clause (original text)
    - Its AI Rewritten Safe Clause
    Returns the path to the generated .docx file.
    """

    tmp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    output_docx_path = tmp_file.name

    doc = Document()

    # --- Title Page ---
    doc.add_heading("High Risk Clause Report", level=0)
    doc.add_paragraph(
        "This document summarizes the high-risk clauses identified in the contract "
        "and provides AI-generated safer rewrites for compliance."
    )

    doc.add_page_break()

    # --- Clauses ---
    for item in analysis_results:
        if item.get("severity", "").lower() == "high":
            original_clause = str(item.get("clause", "")).strip()
            safe_clause = _rewrite_clause_with_llm(item)

            # High risk clause
            doc.add_heading("High Risk Clause:", level=1)
            para = doc.add_paragraph(original_clause)
            para.style.font.size = Pt(11)

            # Rewritten safe clause
            doc.add_heading("Rewritten Safe Clause:", level=1)
            para = doc.add_paragraph(safe_clause)
            para.style.font.size = Pt(11)

            # Separator line
            doc.add_paragraph("-" * 80)

    doc.save(output_docx_path)
    return output_docx_path


def safe_convert_docx_to_pdf(input_path, output_path):
    pythoncom.CoInitialize()
    try:
        output_dir = os.path.dirname(output_path)
        convert(input_path, output_dir)
        generated_pdf = os.path.join(output_dir, os.path.splitext(os.path.basename(input_path))[0] + ".pdf")
        return generated_pdf
    finally:
        pythoncom.CoUninitialize()

def render_download_buttons(analysis_results):
    """
    Generate Word & PDF reports and render Streamlit download buttons.
    """
    try:
        # Step 1: Generate Word file
        docx_path = modify_contract_docx(analysis_results)

        # Step 2: Convert to PDF
        pdf_output_path = os.path.join(tempfile.gettempdir(), "High_Risk_Clause_Report.pdf")
        pdf_path = safe_convert_docx_to_pdf(docx_path, pdf_output_path)

        # --- Word download ---
        with open(docx_path, "rb") as f:
            st.download_button(
                label="⬇️ Download High Risk Clause Report (Word)",
                data=f,
                file_name="High_Risk_Clause_Report.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )

        # --- PDF download ---
        with open(pdf_path, "rb") as f:
            st.download_button(
                label="⬇️ Download High Risk Clause Report (PDF)",
                data=f,
                file_name="High_Risk_Clause_Report.pdf",
                mime="application/pdf",
            )

    except Exception as e:
        st.error(f"❌ Failed to generate reports: {e}")



